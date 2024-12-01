# Libraries
import sqlite3, pandas, matplotlib.pyplot as plt
from matplotlib.widgets import Button
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale
from datetime import datetime

# Connector to the database and to execute the SQL statements the database cursor was created
db = sqlite3.connect ('Resources\\SIEIDB.db')
cur = db.cursor ()

# Validates the access data to the application
def loginv (user, psw):
  cur.execute ('SELECT * FROM UsersSys WHERE users=? AND psw=?', [user, psw])

# Functions to work with the data from the computers that are registered
## Check if the ID is already registered
def id_exist_pc (idpc):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM PC WHERE idpc=?', [idpc])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_pc (idpc, name, model, serial, color, modelmb, colormb, gcn, gcm, cpu, ram, disk, pcso, dp, user, stat, dfa, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO PC (idpc, name, model, serial, color, modelmb, colormb, graphicscardname, graphicscardmodel, cpu, ram, HDDorSDD, so, dp, users, status, dateofarrival, departuredate, observation) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (idpc, name, model, serial, color, modelmb, colormb, gcn, gcm, cpu, ram, disk, pcso, dp, user, stat, dfa, dtd, obs))
  db.commit ()
  db.close ()

## Search engine for data
def search_pc (val, stat, dp):
  query = 'SELECT * FROM PC WHERE'
  params = []

  query += ' (idpc=?'
  params.append(val)

  query += ' OR name LIKE ?)'
  params.append('%'+val+'%')

  query += ' AND status=?'
  params.append(stat)

  if dp !='Todo':
    query += ' AND dp=?'
    params.append(dp)

  cur.execute (query, params)

## Data editor in the database
def edit_pc (rowid, idpc, name, model, serial, color, modelmb, colormb, gcn, gcm, cpu, ram, disk, pcso, dp, user, stat, dom, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE PC SET idpc=:idpc, name=:name, model=:model, serial=:serial, color=:color, modelmb=:modelmb, colormb=:colormb, graphicscardname=:graphicscardname, graphicscardmodel=:graphicscardmodel, cpu=:cpu, ram=:ram, HDDorSDD=:HDDorSDD, so=:so, dp=:dp, users=:users, status=:status, dateofmodification=:dateofmodification, departuredate=:departuredate, observation=:observation WHERE idpc = {idpc}', {'idpc':idpc, 'name':name, 'model':model, 'serial':serial, 'color':color, 'modelmb':modelmb, 'colormb':colormb, 'graphicscardname':gcn, 'graphicscardmodel':gcm, 'cpu':cpu, 'ram':ram, 'HDDorSDD':disk, 'so':pcso, 'dp':dp, 'users':user, 'status':stat, 'dateofmodification':dom, 'departuredate':dtd, 'observation':obs})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_pc (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM PC WHERE rowid = {rowid}')
  db.commit ()
  db.close ()

## Printing a word document with data obtained from the database
def docx_pc (print_data_menu, addressed_to_entry, by_entry, subject_entry, namepc, model, serial, mb, cpu, ram, disk, observation_docx_entry, item_values):
  ### Check that there are enough values
  if len(item_values) < 3:
    messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
    return

  ### Values that the letter document may or may not have
  #### Data to whom the letter is addressed, who sends the letter and the subject
  addressed_to = addressed_to_entry.get().strip ()
  by = by_entry.get().strip ()
  subject = subject_entry.get().strip ()

  #### Assuming the columns to get the data are: Name, Model, Serial, Motherboard, Processor, Memory RAM, Hard disk or solid disk
  pcnamepc = item_values [1] # Name
  pcmodelpc = item_values [2] # Model
  pcserialpc = item_values [3] # Serial
  pcmbpc = item_values [5] # Motherboard
  pccpupc = item_values [9] # Processor
  pcrampc = item_values [10] # Memory RAM
  pcdiskpc = item_values [11] # Hard disk or solid disk

  #### Get the text of the observation
  observation = observation_docx_entry.get('1.0', 'end').strip ()

  ### Create the Word document
  docx = Document()
  section = docx.sections [-1]
  section.page_height = Inches (11)
  section.page_width = Inches (8.5)

  #### Modification of the header to put logos
  header_for_logos = docx.sections[0].header
  hflp = header_for_logos.paragraphs [0]

  ##### Add logo on the left
  tll = hflp.add_run ()
  tll.add_picture ('Resources\\Img\\Logo.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

  ##### Add a space between the images
  hflp.add_run ('                                                                                         ')

  ##### Add logo on the right
  trl = hflp.add_run ()
  trl.add_picture ('Resources\\Img\\Logo2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
  hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  #### Modification of the body for the creation of the document
  ##### Introductory text as a header
  tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
  tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  tihr2 = docx.add_paragraph ('Estado Zulia')
  tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Places the date automatically and also puts it in Spanish
  locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
  current_date = datetime.now().strftime ('%d de %B de %Y')
  tcud = docx.add_paragraph (f'Cabimas, {current_date}')
  tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  ##### Recipient and sender information
  ###### Recipient
  tat = 'Para:'
  details = []

  if addressed_to_entry:
    tat = f'Para: {addressed_to}'

  tatp = docx.add_paragraph (tat)
  tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ###### sender
  tby = 'De:'
  details = []

  if by_entry:
    tby = f'De: {by}'

  tbyp = docx.add_paragraph (tby)
  tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Subject information
  tsu = 'Asunto:'
  details = []

  if subject_entry:
    tsu = f'Asunto: {subject}'

  tsup = docx.add_paragraph (tsu)
  tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Introductory paragraph
  tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
  tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Content for the details of the letter
  tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
  details = []

  ###### Add details according to the selections
  if namepc.get ():
    details.append (f'"{pcnamepc}"')
  if model.get ():
    details.append (f'modelo "{pcmodelpc}"')
  if serial.get ():
    details.append (f'serial "{pcserialpc}"')

  if details:
    tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

  if observation_docx_entry:
    tcd += f" problema: {observation}\n"

  tcdp = docx.add_paragraph (tcd)
  tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  #### Modification of the final part of the body for the creation of the document
  ##### Farewell paragraph
  fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
  fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says who made the letter
  fts = docx.add_paragraph ()
  fts_run = fts.add_run (f'{by_entry.get()}')
  fts_run.bold = True
  fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says which department made the letter
  ftd = docx.add_paragraph ()
  ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
  ftd_run.bold = True
  ftd_run.underline = True
  ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ### Option that opens a window that allows you to decide the name and place where the document will be saved
  file_save = filedialog.asksaveasfilename (defaultextension='.docx', filetypes=[('Documentos de Word', '*.docx'), ('Todos los archivos', '*.*')])

  if file_save:
    docx.save (file_save)
    messagebox.showinfo ('Éxito', 'El documento se ha guardado correctamente.')

  ### Close the window for the generation of the document
  print_data_menu.destroy()

## Statistical graph about computers
### Global variable so that the window does not open more than one
graph_shown = False

def graph_pc ():
  ### Specifying the global variable
  global graph_shown

  ### Conditional that tells me if there is already an open window
  if graph_shown:
    messagebox.showerror ('Nueva ventana', 'Ya existe una ventana abierta')
    return

  ### Connect to the SQLite database
  db = sqlite3.connect ('Resources\\SIEIDB.db')

  ### Query SQL to get data
  cur1 = '''SELECT sum_idpc, status FROM (SELECT COUNT(idpc) AS sum_idpc, status FROM PC GROUP BY status HAVING sum_idpc > 0 UNION SELECT COUNT(idpc) AS total, 'Total General' FROM PC) AS combined_results ORDER BY CASE WHEN status = 'Total General' THEN 0 ELSE 1 END, sum_idpc DESC LIMIT 5'''
  datast = pandas.read_sql (cur1, db)
  cur2 = '''SELECT sum_idpc, dp FROM (SELECT COUNT(idpc) AS sum_idpc, dp FROM PC GROUP BY dp HAVING sum_idpc > 0 UNION SELECT COUNT(idpc) AS total, 'Total General' FROM PC) AS combined_results ORDER BY CASE WHEN dp = 'Total General' THEN 0 ELSE 1 END, sum_idpc DESC LIMIT 5'''
  datadp = pandas.read_sql (cur2, db)

  ### Create a figure and an axis
  fig, ax = plt.subplots ()
  plt.subplots_adjust (bottom=0.2)
  fig.canvas.manager.set_window_title ('Estadistiscas')

  ### Colors
  colors = ['blue', 'green', 'red', 'yellow']

  ### Function to draw the graph
  def draw_graphst (data, title):
    ax.clear ()
    bars = ax.bar (data.status, data.sum_idpc, color=colors[:len(data.status)], label=data.status)
    for bar, label in zip (bars, data.status):
        yval = bar.get_height ()
        ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
        bar.set_label (label)
    ax.legend (title='Estados')
    ax.set_ylabel ('Cantidad')
    ax.set_title (title)
    plt.draw ()

  def draw_graphdp (data, title):
      ax.clear ()
      bars = ax.bar (data.dp, data.sum_idpc, color=colors[:len(data.dp)], label=data.dp)
      for bar, label in zip (bars, data.dp):
          yval = bar.get_height ()
          ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
          bar.set_label (label)
      ax.legend (title='Departamentos')
      ax.set_ylabel ('Cantidad')
      ax.set_title (title)
      plt.draw ()

  ### Initialize the graph
  global current_graph
  current_graph = 1
  draw_graphst (datast, 'PC operativas')
  graph_shown = True

  ### Function to update the graph
  def update_graph (event):
    global current_graph
    if event.inaxes == ax_button_next:
        if current_graph == 1:
            draw_graphdp (datadp, 'PC operativas por departamento')
            current_graph = 2
        else:
            draw_graphst (datast, 'PC operativas')
            current_graph = 1
    elif event.inaxes == ax_button_prev:
        if current_graph == 2:
            draw_graphst (datast, 'PC operativas')
            current_graph = 1
        else:
            draw_graphdp (datadp, 'PC operativas por departamento')
            current_graph = 2

  ### Buttons
  ax_button_prev = plt.axes ([0.1, 0.05, 0.35, 0.075])
  ax_button_next = plt.axes ([0.55, 0.05, 0.35, 0.075])
  button_prev = Button (ax_button_prev, 'Gráfica Anterior')
  button_next = Button (ax_button_next, 'Siguiente Gráfica')

  ### Assign the update function to the buttons
  button_prev.on_clicked (update_graph)
  button_next.on_clicked (update_graph)

  ### Function to reset the variable for me so that I can reopen the window
  def on_close (event):
    global graph_shown
    graph_shown = False
    plt.close (event.canvas.figure)

  ### Event that uses the function to reset the variable that allows only one window to open
  fig.canvas.mpl_connect ('close_event', on_close)

  plt.show()

# Functions to work with the data of the keyboards that are registered
## Check if the ID is already registered
def id_exist_pk (idpk):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM PK WHERE idpk=?', [idpk])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_pk (idpk, name, model, serial, color, dp, user, stat, dfa, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO PK (idpk, name, model, serial, color, dp, users, status, dateofarrival, departuredate, observation) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (idpk, name, model, serial, color, dp, user, stat, dfa, dtd, obs))
  db.commit ()
  db.close ()

## Search engine for data
def search_pk (val, stat, dp):
  query = 'SELECT * FROM PK WHERE'
  params = []

  query += ' (idpk=?'
  params.append(val)

  query += ' OR name LIKE ?)'
  params.append('%'+val+'%')

  query += ' AND status=?'
  params.append(stat)

  if dp !='Todo':
    query += ' AND dp=?'
    params.append(dp)

  cur.execute (query, params)

## Data editor in the database
def edit_pk (rowid, idpk, name, model, serial, color, dp, user, stat, dom, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE PK SET idpk=:idpk, name=:name, model=:model, serial=:serial, color=:color, dp=:dp, users=:users, status=:status, dateofmodification=:dateofmodification, departuredate=:departuredate, observation=:observation WHERE idpk = {idpk}', {'idpk':idpk, 'name':name, 'model':model, 'serial':serial, 'color':color, 'dp':dp, 'users':user, 'status':stat, 'dateofmodification':dom, 'departuredate':dtd, 'observation':obs})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_pk (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM PK WHERE rowid = {rowid}')
  db.commit ()
  db.close ()

## Printing a word document with data obtained from the database
def docx_pk (print_data_menu, addressed_to_entry, by_entry, subject_entry, namepk, model, serial, observation_docx_entry, item_values):
  ### Check that there are enough values
  if len(item_values) < 3:
    messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
    return

  ### Values that the letter document may or may not have
  #### Data to whom the letter is addressed, who sends the letter and the subject
  addressed_to = addressed_to_entry.get().strip ()
  by = by_entry.get().strip ()
  subject = subject_entry.get().strip ()

  #### Assuming the columns to get the data are: Name, Model, Serial
  pknamepk = item_values [1] # Name
  pkmodelpk = item_values [2] # Model
  pkserialpk = item_values [3] # Serial


  #### Get the text of the observation
  observation = observation_docx_entry.get('1.0', 'end').strip ()

  ### Create the Word document
  docx = Document()
  section = docx.sections [-1]
  section.page_height = Inches (11)
  section.page_width = Inches (8.5)

  #### Modification of the header to put logos
  header_for_logos = docx.sections[0].header
  hflp = header_for_logos.paragraphs [0]

  ##### Add logo on the left
  tll = hflp.add_run ()
  tll.add_picture ('Resources\\Img\\Logo.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

  ##### Add a space between the images
  hflp.add_run ('                                                                                         ')

  ##### Add logo on the right
  trl = hflp.add_run ()
  trl.add_picture ('Resources\\Img\\Logo2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
  hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  #### Modification of the body for the creation of the document
  ##### Introductory text as a header
  tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
  tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  tihr2 = docx.add_paragraph ('Estado Zulia')
  tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Places the date automatically and also puts it in Spanish
  locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
  current_date = datetime.now().strftime ('%d de %B de %Y')
  tcud = docx.add_paragraph (f'Cabimas, {current_date}')
  tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  ##### Recipient and sender information
  ###### Recipient
  tat = 'Para:'
  details = []

  if addressed_to_entry:
    tat = f'Para: {addressed_to}'

  tatp = docx.add_paragraph (tat)
  tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ###### sender
  tby = 'De:'
  details = []

  if by_entry:
    tby = f'De: {by}'

  tbyp = docx.add_paragraph (tby)
  tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Subject information
  tsu = 'Asunto:'
  details = []

  if subject_entry:
    tsu = f'Asunto: {subject}'

  tsup = docx.add_paragraph (tsu)
  tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Introductory paragraph
  tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
  tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Content for the details of the letter
  tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
  details = []

  ###### Add details according to the selections
  if namepk.get ():
    details.append (f'"{pknamepk}"')
  if model.get ():
    details.append (f'modelo "{pkmodelpk}"')
  if serial.get ():
    details.append (f'serial "{pkserialpk}"')

  if details:
    tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

  if observation_docx_entry:
    tcd += f" problema: {observation}\n"

  tcdp = docx.add_paragraph (tcd)
  tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  #### Modification of the final part of the body for the creation of the document
  ##### Farewell paragraph
  fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
  fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says who made the letter
  fts = docx.add_paragraph ()
  fts_run = fts.add_run (f'{by_entry.get()}')
  fts_run.bold = True
  fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says which department made the letter
  ftd = docx.add_paragraph ()
  ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
  ftd_run.bold = True
  ftd_run.underline = True
  ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ### Option that opens a window that allows you to decide the name and place where the document will be saved
  file_save = filedialog.asksaveasfilename (defaultextension='.docx', filetypes=[('Documentos de Word', '*.docx'), ('Todos los archivos', '*.*')])

  if file_save:
    docx.save (file_save)
    messagebox.showinfo ('Éxito', 'El documento se ha guardado correctamente.')

  ### Close the window for the generation of the document
  print_data_menu.destroy()

## Statistical graph about the keyboards
### Global variable so that the window does not open more than one
graph_shown = False

def graph_pk ():
  ### Specifying the global variable
  global graph_shown

  ### Conditional that tells me if there is already an open window
  if graph_shown:
    messagebox.showerror ('Nueva ventana', 'Ya existe una ventana abierta')
    return

  ### Connect to the SQLite database
  db = sqlite3.connect ('Resources\\SIEIDB.db')

  ### Query SQL to get data
  cur1 = '''SELECT sum_idpk, status FROM (SELECT COUNT(idpk) AS sum_idpk, status FROM PK GROUP BY status HAVING sum_idpk > 0 UNION SELECT COUNT(idpk) AS total, 'Total General' FROM PK) AS combined_results ORDER BY CASE WHEN status = 'Total General' THEN 0 ELSE 1 END, sum_idpk DESC LIMIT 5'''
  datast = pandas.read_sql (cur1, db)
  cur2 = '''SELECT sum_idpk, dp FROM (SELECT COUNT(idpk) AS sum_idpk, dp FROM PK GROUP BY dp HAVING sum_idpk > 0 UNION SELECT COUNT(idpk) AS total, 'Total General' FROM PK) AS combined_results ORDER BY CASE WHEN dp = 'Total General' THEN 0 ELSE 1 END, sum_idpk DESC LIMIT 5'''
  datadp = pandas.read_sql (cur2, db)

  ### Create a figure and an axis
  fig, ax = plt.subplots ()
  plt.subplots_adjust (bottom=0.2)
  fig.canvas.manager.set_window_title ('Estadistiscas')

  ### Colors
  colors = ['blue', 'green', 'red', 'yellow']

  ### Function to draw the graph
  def draw_graphst (data, title):
    ax.clear ()
    bars = ax.bar (data.status, data.sum_idpk, color=colors[:len(data.status)], label=data.status)
    for bar, label in zip (bars, data.status):
        yval = bar.get_height ()
        ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
        bar.set_label (label)
    ax.legend (title='Estados')
    ax.set_ylabel ('Cantidad')
    ax.set_title (title)
    plt.draw ()

  def draw_graphdp (data, title):
      ax.clear ()
      bars = ax.bar (data.dp, data.sum_idpk, color=colors[:len(data.dp)], label=data.dp)
      for bar, label in zip (bars, data.dp):
          yval = bar.get_height ()
          ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
          bar.set_label (label)
      ax.legend (title='Departamentos')
      ax.set_ylabel ('Cantidad')
      ax.set_title (title)
      plt.draw ()

  ### Initialize the graph
  global current_graph
  current_graph = 1
  draw_graphst (datast, 'Teclados operativos')
  graph_shown = True

  ### Function to update the graph
  def update_graph (event):
    global current_graph
    if event.inaxes == ax_button_next:
        if current_graph == 1:
            draw_graphdp (datadp, 'Teclados operativos por departamento')
            current_graph = 2
        else:
            draw_graphst (datast, 'Teclados operativos')
            current_graph = 1
    elif event.inaxes == ax_button_prev:
        if current_graph == 2:
            draw_graphst (datast, 'Teclados operativos')
            current_graph = 1
        else:
            draw_graphdp (datadp, 'Teclados operativos por departamento')
            current_graph = 2

  ### Buttons
  ax_button_prev = plt.axes ([0.1, 0.05, 0.35, 0.075])
  ax_button_next = plt.axes ([0.55, 0.05, 0.35, 0.075])
  button_prev = Button (ax_button_prev, 'Gráfica Anterior')
  button_next = Button (ax_button_next, 'Siguiente Gráfica')

  ### Assign the update function to the buttons
  button_prev.on_clicked (update_graph)
  button_next.on_clicked (update_graph)

  ### Function to reset the variable for me so that I can reopen the window
  def on_close (event):
    global graph_shown
    graph_shown = False
    plt.close (event.canvas.figure)

  ### Event that uses the function to reset the variable that allows only one window to open
  fig.canvas.mpl_connect ('close_event', on_close)

  plt.show()

# Functions to work with the data of the monitors that are registered
## Check if the ID is already registered
def id_exist_pm (idpm):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM PM WHERE idpm=?', [idpm])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_pm (idpm, name, model, serial, color, tsi, tcp, dp, user, stat, dfa, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO PM (idpm, name, model, serial, color, typescreeninch, typeconnectorport, dp, users, status, dateofarrival, departuredate, observation) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (idpm, name, model, serial, color, tsi, tcp, dp, user, stat, dfa, dtd, obs))
  db.commit ()
  db.close ()

## Search engine for data
def search_pm (val, stat, dp):
  query = 'SELECT * FROM PM WHERE'
  params = []

  query += ' (idpm=?'
  params.append(val)

  query += ' OR name LIKE ?)'
  params.append('%'+val+'%')

  query += ' AND status=?'
  params.append(stat)

  if dp !='Todo':
    query += ' AND dp=?'
    params.append(dp)

  cur.execute (query, params)

## Data editor in the database
def edit_pm (rowid, idpm, name, model, serial, color, tsi, tcp, dp, user, stat, dom, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE PM SET idpm=:idpm, name=:name, model=:model, serial=:serial, color=:color, typescreeninch=:typescreeninch, typeconnectorport=:typeconnectorport, dp=:dp, users=:users, status=:status, dateofmodification=:dateofmodification, departuredate=:departuredate, observation=:observation WHERE idpm = {idpm}', {'idpm':idpm, 'name':name, 'model':model, 'serial':serial, 'color':color, 'typescreeninch':tsi, 'typeconnectorport':tcp, 'dp':dp, 'users':user, 'status':stat, 'dateofmodification':dom, 'departuredate':dtd, 'observation':obs})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_pm (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM PM WHERE rowid = {rowid}')
  db.commit ()
  db.close ()

## Printing a word document with data obtained from the database
def docx_pm (print_data_menu, addressed_to_entry, by_entry, subject_entry, namepm, model, serial, observation_docx_entry, item_values):
  ### Check that there are enough values
  if len(item_values) < 3:
    messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
    return

  ### Values that the letter document may or may not have
  #### Data to whom the letter is addressed, who sends the letter and the subject
  addressed_to = addressed_to_entry.get().strip ()
  by = by_entry.get().strip ()
  subject = subject_entry.get().strip ()

  #### Assuming the columns to get the data are: Name, Model, Serial
  pmnamepm = item_values [1] # Name
  pmmodelpm = item_values [2] # Model
  pmserialpm = item_values [3] # Serial

  #### Get the text of the observation
  observation = observation_docx_entry.get('1.0', 'end').strip ()

  ### Create the Word document
  docx = Document()
  section = docx.sections [-1]
  section.page_height = Inches (11)
  section.page_width = Inches (8.5)

  #### Modification of the header to put logos
  header_for_logos = docx.sections[0].header
  hflp = header_for_logos.paragraphs [0]

  ##### Add logo on the left
  tll = hflp.add_run ()
  tll.add_picture ('Resources\\Img\\Logo.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

  ##### Add a space between the images
  hflp.add_run ('                                                                                         ')

  ##### Add logo on the right
  trl = hflp.add_run ()
  trl.add_picture ('Resources\\Img\\Logo2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
  hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  #### Modification of the body for the creation of the document
  ##### Introductory text as a header
  tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
  tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  tihr2 = docx.add_paragraph ('Estado Zulia')
  tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Places the date automatically and also puts it in Spanish
  locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
  current_date = datetime.now().strftime ('%d de %B de %Y')
  tcud = docx.add_paragraph (f'Cabimas, {current_date}')
  tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  ##### Recipient and sender information
  ###### Recipient
  tat = 'Para:'
  details = []

  if addressed_to_entry:
    tat = f'Para: {addressed_to}'

  tatp = docx.add_paragraph (tat)
  tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ###### sender
  tby = 'De:'
  details = []

  if by_entry:
    tby = f'De: {by}'

  tbyp = docx.add_paragraph (tby)
  tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Subject information
  tsu = 'Asunto:'
  details = []

  if subject_entry:
    tsu = f'Asunto: {subject}'

  tsup = docx.add_paragraph (tsu)
  tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Introductory paragraph
  tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
  tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Content for the details of the letter
  tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
  details = []

  ###### Add details according to the selections
  if namepm.get ():
    details.append (f'"{pmnamepm}"')
  if model.get ():
    details.append (f'modelo "{pmmodelpm}"')
  if serial.get ():
    details.append (f'serial "{pmserialpm}"')

  if details:
    tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

  if observation_docx_entry:
    tcd += f" problema: {observation}\n"

  tcdp = docx.add_paragraph (tcd)
  tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  #### Modification of the final part of the body for the creation of the document
  ##### Farewell paragraph
  fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
  fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says who made the letter
  fts = docx.add_paragraph ()
  fts_run = fts.add_run (f'{by_entry.get()}')
  fts_run.bold = True
  fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says which department made the letter
  ftd = docx.add_paragraph ()
  ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
  ftd_run.bold = True
  ftd_run.underline = True
  ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ### Option that opens a window that allows you to decide the name and place where the document will be saved
  file_save = filedialog.asksaveasfilename (defaultextension='.docx', filetypes=[('Documentos de Word', '*.docx'), ('Todos los archivos', '*.*')])

  if file_save:
    docx.save (file_save)
    messagebox.showinfo ('Éxito', 'El documento se ha guardado correctamente.')

  ### Close the window for the generation of the document
  print_data_menu.destroy()

## Statistical graph about the monitors
### Global variable so that the window does not open more than one
graph_shown = False

def graph_pm ():
  ### Specifying the global variable
  global graph_shown

  ### Conditional that tells me if there is already an open window
  if graph_shown:
    messagebox.showerror ('Nueva ventana', 'Ya existe una ventana abierta')
    return

  ### Connect to the SQLite database
  db = sqlite3.connect ('Resources\\SIEIDB.db')

  ### Query SQL to get data
  cur1 = '''SELECT sum_idpm, status FROM (SELECT COUNT(idpm) AS sum_idpm, status FROM PM GROUP BY status HAVING sum_idpm > 0 UNION SELECT COUNT(idpm) AS total, 'Total General' FROM PM) AS combined_results ORDER BY CASE WHEN status = 'Total General' THEN 0 ELSE 1 END, sum_idpm DESC LIMIT 5'''
  datast = pandas.read_sql (cur1, db)
  cur2 = '''SELECT sum_idpm, dp FROM (SELECT COUNT(idpm) AS sum_idpm, dp FROM PM GROUP BY dp HAVING sum_idpm > 0 UNION SELECT COUNT(idpm) AS total, 'Total General' FROM PM) AS combined_results ORDER BY CASE WHEN dp = 'Total General' THEN 0 ELSE 1 END, sum_idpm DESC LIMIT 5'''
  datadp = pandas.read_sql (cur2, db)

  ### Create a figure and an axis
  fig, ax = plt.subplots ()
  plt.subplots_adjust (bottom=0.2)
  fig.canvas.manager.set_window_title ('Estadistiscas')

  ### Colors
  colors = ['blue', 'green', 'red', 'yellow']

  ### Function to draw the graph
  def draw_graphst (data, title):
    ax.clear ()
    bars = ax.bar (data.status, data.sum_idpm, color=colors[:len(data.status)], label=data.status)
    for bar, label in zip (bars, data.status):
        yval = bar.get_height ()
        ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
        bar.set_label (label)
    ax.legend (title='Estados')
    ax.set_ylabel ('Cantidad')
    ax.set_title (title)
    plt.draw ()

  def draw_graphdp (data, title):
      ax.clear ()
      bars = ax.bar (data.dp, data.sum_idpm, color=colors[:len(data.dp)], label=data.dp)
      for bar, label in zip (bars, data.dp):
          yval = bar.get_height ()
          ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
          bar.set_label (label)
      ax.legend (title='Departamentos')
      ax.set_ylabel ('Cantidad')
      ax.set_title (title)
      plt.draw ()

  ### Initialize the graph
  global current_graph
  current_graph = 1
  draw_graphst (datast, 'Monitores operativos')
  graph_shown = True

  ### Function to update the graph
  def update_graph (event):
    global current_graph
    if event.inaxes == ax_button_next:
        if current_graph == 1:
            draw_graphdp (datadp, 'Monitores operativos por departamento')
            current_graph = 2
        else:
            draw_graphst (datast, 'Monitores operativos')
            current_graph = 1
    elif event.inaxes == ax_button_prev:
        if current_graph == 2:
            draw_graphst (datast, 'Monitores operativos')
            current_graph = 1
        else:
            draw_graphdp (datadp, 'Monitores operativos por departamento')
            current_graph = 2

  ### Buttons
  ax_button_prev = plt.axes ([0.1, 0.05, 0.35, 0.075])
  ax_button_next = plt.axes ([0.55, 0.05, 0.35, 0.075])
  button_prev = Button (ax_button_prev, 'Gráfica Anterior')
  button_next = Button (ax_button_next, 'Siguiente Gráfica')

  ### Assign the update function to the buttons
  button_prev.on_clicked (update_graph)
  button_next.on_clicked (update_graph)

  ### Function to reset the variable for me so that I can reopen the window
  def on_close (event):
    global graph_shown
    graph_shown = False
    plt.close (event.canvas.figure)

  ### Event that uses the function to reset the variable that allows only one window to open
  fig.canvas.mpl_connect ('close_event', on_close)

  plt.show()

# Functions to work with the data of the mouses that are registered
## Check if the ID is already registered
def id_exist_pmo (idpmo):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM PMO WHERE idpmo=?', [idpmo])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_pmo (idpmo, name, model, serial, color, dp, user, stat, dfa, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO PMO (idpmo, name, model, serial, color, dp, users, status, dateofarrival, departuredate, observation) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (idpmo, name, model, serial, color, dp, user, stat, dfa, dtd, obs))
  db.commit ()
  db.close ()

## Search engine for data
def search_pmo (val, stat, dp):
  query = 'SELECT * FROM PMO WHERE'
  params = []

  query += ' (idpmo=?'
  params.append(val)

  query += ' OR name LIKE ?)'
  params.append('%'+val+'%')

  query += ' AND status=?'
  params.append(stat)

  if dp !='Todo':
    query += ' AND dp=?'
    params.append(dp)

  cur.execute (query, params)

## Data editor in the database
def edit_pmo (rowid, idpmo, name, model, serial, color, dp, user, stat, dom, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE PMO SET idpmo=:idpmo, name=:name, model=:model, serial=:serial, color=:color, dp=:dp, users=:users, status=:status, dateofmodification=:dateofmodification, departuredate=:departuredate, observation=:observation WHERE idpmo = {idpmo}', {'idpmo':idpmo, 'name':name, 'model':model, 'serial':serial, 'color':color, 'dp':dp, 'users':user, 'status':stat, 'dateofmodification':dom, 'departuredate':dtd, 'observation':obs})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_pmo (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM PMO WHERE rowid = {rowid}')
  db.commit ()
  db.close ()

## Printing a word document with data obtained from the database
def docx_pmo (print_data_menu, addressed_to_entry, by_entry, subject_entry, namepmo, model, serial, observation_docx_entry, item_values):
  ### Check that there are enough values
  if len(item_values) < 3:
    messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
    return

  ### Values that the letter document may or may not have
  #### Data to whom the letter is addressed, who sends the letter and the subject
  addressed_to = addressed_to_entry.get().strip ()
  by = by_entry.get().strip ()
  subject = subject_entry.get().strip ()

  #### Assuming the columns to get the data are: Name, Model, Serial
  pmonamepmo = item_values [1] # Name
  pmomodelpmo = item_values [2] # Model
  pmoserialpmo = item_values [3] # Serial


  #### Get the text of the observation
  observation = observation_docx_entry.get('1.0', 'end').strip ()

  ### Create the Word document
  docx = Document()
  section = docx.sections [-1]
  section.page_height = Inches (11)
  section.page_width = Inches (8.5)

  #### Modification of the header to put logos
  header_for_logos = docx.sections[0].header
  hflp = header_for_logos.paragraphs [0]

  ##### Add logo on the left
  tll = hflp.add_run ()
  tll.add_picture ('Resources\\Img\\Logo.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

  ##### Add a space between the images
  hflp.add_run ('                                                                                         ')

  ##### Add logo on the right
  trl = hflp.add_run ()
  trl.add_picture ('Resources\\Img\\Logo2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
  hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  #### Modification of the body for the creation of the document
  ##### Introductory text as a header
  tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
  tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  tihr2 = docx.add_paragraph ('Estado Zulia')
  tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Places the date automatically and also puts it in Spanish
  locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
  current_date = datetime.now().strftime ('%d de %B de %Y')
  tcud = docx.add_paragraph (f'Cabimas, {current_date}')
  tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  ##### Recipient and sender information
  ###### Recipient
  tat = 'Para:'
  details = []

  if addressed_to_entry:
    tat = f'Para: {addressed_to}'

  tatp = docx.add_paragraph (tat)
  tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ###### sender
  tby = 'De:'
  details = []

  if by_entry:
    tby = f'De: {by}'

  tbyp = docx.add_paragraph (tby)
  tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Subject information
  tsu = 'Asunto:'
  details = []

  if subject_entry:
    tsu = f'Asunto: {subject}'

  tsup = docx.add_paragraph (tsu)
  tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Introductory paragraph
  tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
  tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Content for the details of the letter
  tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
  details = []

  ###### Add details according to the selections
  if namepmo.get ():
    details.append (f'"{pmonamepmo}"')
  if model.get ():
    details.append (f'modelo "{pmomodelpmo}"')
  if serial.get ():
    details.append (f'serial "{pmoserialpmo}"')

  if details:
    tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

  if observation_docx_entry:
    tcd += f" problema: {observation}\n"

  tcdp = docx.add_paragraph (tcd)
  tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  #### Modification of the final part of the body for the creation of the document
  ##### Farewell paragraph
  fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
  fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says who made the letter
  fts = docx.add_paragraph ()
  fts_run = fts.add_run (f'{by_entry.get()}')
  fts_run.bold = True
  fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says which department made the letter
  ftd = docx.add_paragraph ()
  ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
  ftd_run.bold = True
  ftd_run.underline = True
  ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ### Option that opens a window that allows you to decide the name and place where the document will be saved
  file_save = filedialog.asksaveasfilename (defaultextension='.docx', filetypes=[('Documentos de Word', '*.docx'), ('Todos los archivos', '*.*')])

  if file_save:
    docx.save (file_save)
    messagebox.showinfo ('Éxito', 'El documento se ha guardado correctamente.')

  ### Close the window for the generation of the document
  print_data_menu.destroy()

## Statistical graph about mouses
### Global variable so that the window does not open more than one
graph_shown = False

def graph_pmo ():
  ### Specifying the global variable
  global graph_shown

  ### Conditional that tells me if there is already an open window
  if graph_shown:
    messagebox.showerror ('Nueva ventana', 'Ya existe una ventana abierta')
    return

  ### Connect to the SQLite database
  db = sqlite3.connect ('Resources\\SIEIDB.db')

  ### Query SQL to get data
  cur1 = '''SELECT sum_idpmo, status FROM (SELECT COUNT(idpmo) AS sum_idpmo, status FROM PMO GROUP BY status HAVING sum_idpmo > 0 UNION SELECT COUNT(idpmo) AS total, 'Total General' FROM PMO) AS combined_results ORDER BY CASE WHEN status = 'Total General' THEN 0 ELSE 1 END, sum_idpmo DESC LIMIT 5'''
  datast = pandas.read_sql (cur1, db)
  cur2 = '''SELECT sum_idpmo, dp FROM (SELECT COUNT(idpmo) AS sum_idpmo, dp FROM PMO GROUP BY dp HAVING sum_idpmo > 0 UNION SELECT COUNT(idpmo) AS total, 'Total General' FROM PMO) AS combined_results ORDER BY CASE WHEN dp = 'Total General' THEN 0 ELSE 1 END, sum_idpmo DESC LIMIT 5'''
  datadp = pandas.read_sql (cur2, db)

  ### Create a figure and an axis
  fig, ax = plt.subplots ()
  plt.subplots_adjust (bottom=0.2)
  fig.canvas.manager.set_window_title ('Estadistiscas')

  ### Colors
  colors = ['blue', 'green', 'red', 'yellow']

  ### Function to draw the graph
  def draw_graphst (data, title):
    ax.clear ()
    bars = ax.bar (data.status, data.sum_idpmo, color=colors[:len(data.status)], label=data.status)
    for bar, label in zip (bars, data.status):
        yval = bar.get_height ()
        ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
        bar.set_label (label)
    ax.legend (title='Estados')
    ax.set_ylabel ('Cantidad')
    ax.set_title (title)
    plt.draw ()

  def draw_graphdp (data, title):
      ax.clear ()
      bars = ax.bar (data.dp, data.sum_idpmo, color=colors[:len(data.dp)], label=data.dp)
      for bar, label in zip (bars, data.dp):
          yval = bar.get_height ()
          ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
          bar.set_label (label)
      ax.legend (title='Departamentos')
      ax.set_ylabel ('Cantidad')
      ax.set_title (title)
      plt.draw ()

  ### Initialize the graph
  global current_graph
  current_graph = 1
  draw_graphst (datast, 'Mouses operativos')
  graph_shown = True

  ### Function to update the graph
  def update_graph (event):
    global current_graph
    if event.inaxes == ax_button_next:
        if current_graph == 1:
            draw_graphdp (datadp, 'Mouses operativos por departamento')
            current_graph = 2
        else:
            draw_graphst (datast, 'Mouses operativos')
            current_graph = 1
    elif event.inaxes == ax_button_prev:
        if current_graph == 2:
            draw_graphst (datast, 'Mouses operativos')
            current_graph = 1
        else:
            draw_graphdp (datadp, 'Mouses operativos por departamento')
            current_graph = 2

  ### Buttons
  ax_button_prev = plt.axes ([0.1, 0.05, 0.35, 0.075])
  ax_button_next = plt.axes ([0.55, 0.05, 0.35, 0.075])
  button_prev = Button (ax_button_prev, 'Gráfica Anterior')
  button_next = Button (ax_button_next, 'Siguiente Gráfica')

  ### Assign the update function to the buttons
  button_prev.on_clicked (update_graph)
  button_next.on_clicked (update_graph)

  ### Function to reset the variable for me so that I can reopen the window
  def on_close (event):
    global graph_shown
    graph_shown = False
    plt.close (event.canvas.figure)

  ### Event that uses the function to reset the variable that allows only one window to open
  fig.canvas.mpl_connect ('close_event', on_close)

  plt.show()

# Functions to work with the data from the printers that are registered
## Check if the ID is already registered
def id_exist_pp (idpp):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM PP WHERE idpp=?', [idpp])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_pp (idpp, name, model, serial, color, tp, dp, user, stat, dfa, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO PP (idpp, name, model, serial, color, typeprinting, dp, users, status, dateofarrival, departuredate, observation) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (idpp, name, model, serial, color, tp, dp, user, stat, dfa, dtd, obs))
  db.commit ()
  db.close ()

## Search engine for data
def search_pp (val, stat, dp):
  query = 'SELECT * FROM PP WHERE'
  params = []

  query += ' (idpp=?'
  params.append(val)

  query += ' OR name LIKE ?)'
  params.append('%'+val+'%')

  query += ' AND status=?'
  params.append(stat)

  if dp !='Todo':
    query += ' AND dp=?'
    params.append(dp)

  cur.execute (query, params)

## Data editor in the database
def edit_pp (rowid, idpp, name, model, serial, color, tp, dp, user, stat, dom, dtd, obs):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE PP SET idpp=:idpp, name=:name, model=:model, serial=:serial, color=:color, typeprinting=:typeprinting, dp=:dp, users=:users, status=:status, dateofmodification=:dateofmodification, departuredate=:departuredate, observation=:observation WHERE idpp = {idpp}', {'idpp':idpp, 'name':name, 'model':model, 'serial':serial, 'color':color, 'typeprinting':tp, 'dp':dp, 'users':user, 'status':stat, 'dateofmodification':dom, 'departuredate':dtd, 'observation':obs})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_pp (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM PP WHERE rowid = {rowid}')
  db.commit ()
  db.close ()

## Printing a word document with data obtained from the database
def docx_pp (print_data_menu, addressed_to_entry, by_entry, subject_entry, namepp, model, serial, typeprinting, observation_docx_entry, item_values):
  ### Check that there are enough values
  if len(item_values) < 3:
    messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
    return

  ### Values that the letter document may or may not have
  #### Data to whom the letter is addressed, who sends the letter and the subject
  addressed_to = addressed_to_entry.get().strip ()
  by = by_entry.get().strip ()
  subject = subject_entry.get().strip ()

  #### Assuming the columns to get the data are: Name, Model, Serial, Type printing
  ppnamepp = item_values [1] # Name
  ppmodelpp = item_values [2] # Model
  ppserialpp = item_values [3] # Serial
  pptypeprintingpp = item_values [3] # Type printing

  #### Get the text of the observation
  observation = observation_docx_entry.get('1.0', 'end').strip ()

  ### Create the Word document
  docx = Document()
  section = docx.sections [-1]
  section.page_height = Inches (11)
  section.page_width = Inches (8.5)

  #### Modification of the header to put logos
  header_for_logos = docx.sections[0].header
  hflp = header_for_logos.paragraphs [0]

  ##### Add logo on the left
  tll = hflp.add_run ()
  tll.add_picture ('Resources\\Img\\Logo.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

  ##### Add a space between the images
  hflp.add_run ('                                                                                         ')

  ##### Add logo on the right
  trl = hflp.add_run ()
  trl.add_picture ('Resources\\Img\\Logo2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
  hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  #### Modification of the body for the creation of the document
  ##### Introductory text as a header
  tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
  tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  tihr2 = docx.add_paragraph ('Estado Zulia')
  tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Places the date automatically and also puts it in Spanish
  locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
  current_date = datetime.now().strftime ('%d de %B de %Y')
  tcud = docx.add_paragraph (f'Cabimas, {current_date}')
  tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  ##### Recipient and sender information
  ###### Recipient
  tat = 'Para:'
  details = []

  if addressed_to_entry:
    tat = f'Para: {addressed_to}'

  tatp = docx.add_paragraph (tat)
  tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ###### sender
  tby = 'De:'
  details = []

  if by_entry:
    tby = f'De: {by}'

  tbyp = docx.add_paragraph (tby)
  tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Subject information
  tsu = 'Asunto:'
  details = []

  if subject_entry:
    tsu = f'Asunto: {subject}'

  tsup = docx.add_paragraph (tsu)
  tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Introductory paragraph
  tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
  tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  ##### Content for the details of the letter
  tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
  details = []

  ###### Add details according to the selections
  if namepp.get ():
    details.append (f'"{ppnamepp}"')
  if model.get ():
    details.append (f'modelo "{ppmodelpp}"')
  if serial.get ():
    details.append (f'serial "{ppserialpp}"')

  if details:
    tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

  if observation_docx_entry:
    tcd += f" problema: {observation}\n"

  tcdp = docx.add_paragraph (tcd)
  tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  #### Modification of the final part of the body for the creation of the document
  ##### Farewell paragraph
  fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
  fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says who made the letter
  fts = docx.add_paragraph ()
  fts_run = fts.add_run (f'{by_entry.get()}')
  fts_run.bold = True
  fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ##### Blank line breaks
  docx.add_paragraph ()
  docx.add_paragraph ()

  ##### Final text that says which department made the letter
  ftd = docx.add_paragraph ()
  ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
  ftd_run.bold = True
  ftd_run.underline = True
  ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  ### Option that opens a window that allows you to decide the name and place where the document will be saved
  file_save = filedialog.asksaveasfilename (defaultextension='.docx', filetypes=[('Documentos de Word', '*.docx'), ('Todos los archivos', '*.*')])

  if file_save:
    docx.save (file_save)
    messagebox.showinfo ('Éxito', 'El documento se ha guardado correctamente.')

  ### Close the window for the generation of the document
  print_data_menu.destroy()

## Statistical graph about printers
### Global variable so that the window does not open more than one
graph_shown = False

def graph_pp ():
  ### Specifying the global variable
  global graph_shown

  ### Conditional that tells me if there is already an open window
  if graph_shown:
    messagebox.showerror ('Nueva ventana', 'Ya existe una ventana abierta')
    return

  ### Connect to the SQLite database
  db = sqlite3.connect ('Resources\\SIEIDB.db')

  ### Query SQL to get data
  cur1 = '''SELECT sum_idpp, status FROM (SELECT COUNT(idpp) AS sum_idpp, status FROM PP GROUP BY status HAVING sum_idpp > 0 UNION SELECT COUNT(idpp) AS total, 'Total General' FROM PP) AS combined_results ORDER BY CASE WHEN status = 'Total General' THEN 0 ELSE 1 END, sum_idpp DESC LIMIT 5'''
  datast = pandas.read_sql (cur1, db)
  cur2 = '''SELECT sum_idpp, dp FROM (SELECT COUNT(idpp) AS sum_idpp, dp FROM PP GROUP BY dp HAVING sum_idpp > 0 UNION SELECT COUNT(idpp) AS total, 'Total General' FROM PP) AS combined_results ORDER BY CASE WHEN dp = 'Total General' THEN 0 ELSE 1 END, sum_idpp DESC LIMIT 5'''
  datadp = pandas.read_sql (cur2, db)

  ### Create a figure and an axis
  fig, ax = plt.subplots ()
  plt.subplots_adjust (bottom=0.2)
  fig.canvas.manager.set_window_title ('Estadistiscas')

  ### Colors
  colors = ['blue', 'green', 'red', 'yellow']

  ### Function to draw the graph
  def draw_graphst (data, title):
    ax.clear ()
    bars = ax.bar (data.status, data.sum_idpp, color=colors[:len(data.status)], label=data.status)
    for bar, label in zip (bars, data.status):
        yval = bar.get_height ()
        ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
        bar.set_label (label)
    ax.legend (title='Estados')
    ax.set_ylabel ('Cantidad')
    ax.set_title (title)
    plt.draw ()

  def draw_graphdp (data, title):
      ax.clear ()
      bars = ax.bar (data.dp, data.sum_idpp, color=colors[:len(data.dp)], label=data.dp)
      for bar, label in zip (bars, data.dp):
          yval = bar.get_height ()
          ax.annotate (yval, (bar.get_x() + bar.get_width() / 2, yval), ha='center', va='bottom')
          bar.set_label (label)
      ax.legend (title='Departamentos')
      ax.set_ylabel ('Cantidad')
      ax.set_title (title)
      plt.draw ()

  ### Initialize the graph
  global current_graph
  current_graph = 1
  draw_graphst (datast, 'Impresoras operativas')
  graph_shown = True

  ### Function to update the graph
  def update_graph (event):
    global current_graph
    if event.inaxes == ax_button_next:
        if current_graph == 1:
            draw_graphdp (datadp, 'Impresoras operativas por departamento')
            current_graph = 2
        else:
            draw_graphst (datast, 'Impresoras operativas')
            current_graph = 1
    elif event.inaxes == ax_button_prev:
        if current_graph == 2:
            draw_graphst (datast, 'Impresoras operativas')
            current_graph = 1
        else:
            draw_graphdp (datadp, 'Impresoras operativas por departamento')
            current_graph = 2

  ### Buttons
  ax_button_prev = plt.axes ([0.1, 0.05, 0.35, 0.075])
  ax_button_next = plt.axes ([0.55, 0.05, 0.35, 0.075])
  button_prev = Button (ax_button_prev, 'Gráfica Anterior')
  button_next = Button (ax_button_next, 'Siguiente Gráfica')

  ### Assign the update function to the buttons
  button_prev.on_clicked (update_graph)
  button_next.on_clicked (update_graph)

  ### Function to reset the variable for me so that I can reopen the window
  def on_close (event):
    global graph_shown
    graph_shown = False
    plt.close (event.canvas.figure)

  ### Event that uses the function to reset the variable that allows only one window to open
  fig.canvas.mpl_connect ('close_event', on_close)

  plt.show()

# Functions to work with the data of users who are registered
## Check if the ID is already registered
def id_exist_users (idus):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('SELECT COUNT(*) FROM UsersSys WHERE idus=?', [idus])
  result = cur.fetchone ()
  db.close ()
  return result [0] > 0

## Inserting data to the database
def insert_users (idus, user, psw, firstnameperson, lastnameperson, idcardperson):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute ('INSERT INTO UsersSys (idus, users, psw, firstnameperson, lastnameperson, idcardperson) VALUES(?, ?, ? ,? ,? ,?)', (idus, user, psw, firstnameperson, lastnameperson, idcardperson))
  db.commit ()
  db.close ()

## Search engine for data
def search_users ():
  cur.execute ('SELECT * FROM UsersSys')

## Data editor in the database
def edit_users (rowid, idus, user, psw, firstnameperson, lastnameperson, idcardperson):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'UPDATE UsersSys SET idus=:idus, users=:user, psw=:psw, firstnameperson=:firstnameperson, lastnameperson=:lastnameperson, idcardperson=:idcardperson WHERE idus = {idus}', {'idus':idus, 'user':user, 'psw':psw, 'firstnameperson':firstnameperson, 'lastnameperson':lastnameperson, 'idcardperson':idcardperson})
  db.commit ()
  db.close ()

## Deletion of data from the database
def del_users (rowid):
  db = sqlite3.connect ('Resources\\SIEIDB.db')
  cur = db.cursor ()
  cur.execute (f'DELETE FROM UsersSys WHERE rowid = {rowid}')
  db.commit ()
  db.close ()
