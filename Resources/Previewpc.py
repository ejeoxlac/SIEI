import tkinter as tk
from tkinter import messagebox, ttk
import os
import pymupdf
import threading
import time
from docx2pdf import convert
import locale
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import pythoncom

def docx_pc_pv(addressed_to_entry, by_entry, subject_entry, namepc, model, serial, mb, cpu, ram, disk, observation_docx_entry, item_values):
    ### Check that there are enough values
    if len(item_values) < 3:
        messagebox.showerror ('Error', 'El registro seleccionado no tiene suficientes datos.')
        return

    ### Values that the letter document may or may not have
    #### Data to whom the letter is addressed, who sends the letter and the subject
    addressed_to = addressed_to_entry.get().strip ()
    by = by_entry.get().strip ()
    subject = subject_entry.get().strip ()

    #### Assuming the columns to get the data are: Name, Model, Serial, Motherboard, Processor, Memory RAM, Hard disk or solid disk
    pcnamepc = item_values [1] # Name
    pcmodelpc = item_values [2] # Model
    pcserialpc = item_values [3] # Serial
    pcmbpc = item_values [5] # Motherboard
    pccpupc = item_values [9] # Processor
    pcrampc = item_values [10] # Memory RAM
    pcdiskpc = item_values [11] # Hard disk or solid disk

    #### Get the text of the observation
    observation = observation_docx_entry.get('1.0', 'end').strip ()

    # Crear la carpeta "temp" si no existe
    carpeta_temp = "temp"
    if not os.path.exists(carpeta_temp):
        os.makedirs(carpeta_temp)

    ### Create the Word document
    docx = Document()
    section = docx.sections [-1]
    section.page_height = Inches (11)
    section.page_width = Inches (8.5)

    #### Modification of the header to put logos
    header_for_logos = docx.sections[0].header
    hflp = header_for_logos.paragraphs [0]

    ##### Add logo on the left
    tll = hflp.add_run ()
    tll.add_picture ('Resources\\Img\\Logo_for_SIEI_1.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen

    ##### Add a space between the images
    hflp.add_run ('                                                                                         ')

    ##### Add logo on the right
    trl = hflp.add_run ()
    trl.add_picture ('Resources\\Img\\Logo_for_SIEI_2.png', width=Inches(1.5))  # Ajusta el tamaño de la imagen
    hflp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #### Modification of the body for the creation of the document
    ##### Introductory text as a header
    tihr1= docx.add_paragraph ('Republica borivariana de venezuela')
    tihr1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    tihr2 = docx.add_paragraph ('Estado Zulia')
    tihr2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ##### Places the date automatically and also puts it in Spanish
    locale.setlocale (locale.LC_TIME, 'es_ES.UTF-8')  # Switch to 'es_ES' if you have problems with 'es_ES.UTF-8'
    current_date = datetime.now().strftime ('%d de %B de %Y')
    tcud = docx.add_paragraph (f'Cabimas, {current_date}')
    tcud.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    ##### Recipient and sender information
    ###### Recipient
    tat = 'Para:'
    details = []

    if addressed_to_entry:
        tat = f'Para: {addressed_to}'

    tatp = docx.add_paragraph (tat)
    tatp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    ###### sender
    tby = 'De:'
    details = []

    if by_entry:
        tby = f'De: {by}'

    tbyp = docx.add_paragraph (tby)
    tbyp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    ##### Subject information
    tsu = 'Asunto:'
    details = []

    if subject_entry:
        tsu = f'Asunto: {subject}'

    tsup = docx.add_paragraph (tsu)
    tsup.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    ##### Introductory paragraph
    tinp = docx.add_paragraph ('Reciba usted un cordial saludo')
    tinp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    ##### Content for the details of the letter
    tcd = 'El equipo de la marca, modelo, serial, está presentando el siguiente problema, sus caracteriticas son:\n'
    details = []

    ###### Add details according to the selections
    if namepc.get ():
        details.append (f'"{pcnamepc}"')
    if model.get ():
        details.append (f'modelo "{pcmodelpc}"')
    if serial.get ():
        details.append (f'serial "{pcserialpc}"')

    if details:
        tcd = f'El equipo de la marca {', '.join(details)} está presentando el siguiente'

    if observation_docx_entry:
        tcd += f" problema: {observation}\n"

    tcdp = docx.add_paragraph (tcd)
    tcdp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #### Modification of the final part of the body for the creation of the document
    ##### Farewell paragraph
    fwp = docx.add_paragraph ('Sin más que comentar, me despido de usted')
    fwp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ##### Blank line breaks
    docx.add_paragraph ()
    docx.add_paragraph ()

    ##### Final text that says who made the letter
    fts = docx.add_paragraph ()
    fts_run = fts.add_run (f'{by_entry.get()}')
    fts_run.bold = True
    fts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ##### Blank line breaks
    docx.add_paragraph ()
    docx.add_paragraph ()

    ##### Final text that says which department made the letter
    ftd = docx.add_paragraph ()
    ftd_run = ftd.add_run ('De parte de la Dirección de Informática')
    ftd_run.bold = True
    ftd_run.underline = True
    ftd.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Definir la ruta del archivo DOCX
    doc_path = os.path.join(carpeta_temp, "temp_document_pc.docx")
    
    # Guardar el documento
    try:
        docx.save(doc_path)
        print(f"Documento guardado en: {doc_path}")  # Mensaje de depuración
    except Exception as e:
        print(f"Error al guardar el documento: {e}")  # Mensaje de error
        return None
    
    return doc_path

def convertir_a_pdf(addressed_to_entry, by_entry, subject_entry, namepc, model, serial, mb, cpu, ram, disk, observation_docx_entry, item_values):
    archivo_docx = docx_pc_pv (addressed_to_entry, by_entry, subject_entry, namepc, model, serial, mb, cpu, ram, disk, observation_docx_entry, item_values)  # Generar el documento DOCX automáticamente

    if archivo_docx:
        ventana_progreso = tk.Toplevel()
        ventana_progreso.title("Progreso de Conversión")
        ventana_progreso.geometry("300x100")

        # Crear una barra de progreso
        barra_progreso = ttk.Progressbar(ventana_progreso, orient="horizontal", length=250, mode="indeterminate")
        barra_progreso.pack(pady=20)

        barra_progreso.start(10)  # Iniciar la barra de progreso

        # Iniciar el hilo para realizar la conversión
        threading.Thread(target=realizar_conversion, args=(archivo_docx, ventana_progreso)).start()

    if archivo_docx is None:
        messagebox.showerror("Error", "No se pudo generar el documento DOCX.")
        return  # Salir si no se generó el DOCX

def realizar_conversion(archivo_docx, ventana_progreso):
    pythoncom.CoInitialize()  # Inicializa COM en este hilo
    try:
        # Definir la ruta del archivo PDF
        pdf_path = os.path.join("temp", 'temppc.pdf')  # Guardar como temp.pdf en la carpeta temp
        convert(archivo_docx, pdf_path)  # Especificar la ruta de salida
        print(f"PDF guardado en: {pdf_path}")  # Mensaje de depuración

        # Mostrar el PDF en una nueva ventana
        mostrar_pdf(pdf_path, archivo_docx)
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")  # Mensaje de error
        messagebox.showerror("Error", f"Error al convertir a PDF: {e}")  # Mostrar mensaje de error
    finally:
        pythoncom.CoUninitialize()  # Desinicializa COM
        ventana_progreso.destroy()  # Cerrar la ventana de progreso

def mostrar_pdf(pdf_path, docx_path):
    # Crear una nueva ventana para mostrar el PDF
    ventana_pdf = tk.Toplevel()
    ventana_pdf.title("Visualizador de PDF")
    ventana_pdf.geometry("960x600")  # Establecer el tamaño de la ventana a 960x600

    # Crear un marco para el lienzo y el scrollbar
    marco = tk.Frame(ventana_pdf)
    marco.pack(pady=20)

    # Crear un lienzo para mostrar el PDF
    lienzo = tk.Canvas(marco, width=600, height=800)
    lienzo.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Crear un scrollbar
    scrollbar = tk.Scrollbar(marco, orient="vertical", command=lienzo.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    lienzo.configure(yscrollcommand=scrollbar.set)

    # Abrir el PDF
    doc = pymupdf.open(pdf_path)
    temp_images = []  # Lista para almacenar las rutas de las imágenes temporales
    for page_num in range(len(doc)):
        pagina = doc.load_page(page_num)  # Cargar cada página
        pix = pagina.get_pixmap()  # Renderizar la página como imagen

        # Guardar la imagen temporalmente
        img_path = os.path.join("temp", f"temp_image_pc_{page_num}.png")  # Guardar en la carpeta temp
        pix.save(img_path)
        temp_images.append(img_path)  # Agregar la ruta a la lista

        # Cargar la imagen en el lienzo
        img = tk.PhotoImage(file=img_path)
        lienzo.create_image(0, page_num * 800, anchor=tk.NW, image=img)
        lienzo.image = img  # Mantener una referencia a la imagen

    # Ajustar el tamaño del lienzo para que se pueda desplazar
    lienzo.config(scrollregion=lienzo.bbox("all"))

    # Función para limpiar los archivos temporales
    def cleanup_temp_files():
        # Eliminar las imágenes temporales
        for img_file in temp_images:
            if os.path.exists(img_file):
                os.remove(img_file)
        # Eliminar el archivo PDF temporal
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        # Eliminar el archivo DOCX temporal
        if os.path.exists(docx_path):
            os.remove(docx_path)
        # Cerrar la ventana
        ventana_pdf.destroy()

    # Vincular la función de limpieza al evento de cierre de la ventana
    ventana_pdf.protocol("WM_DELETE_WINDOW", cleanup_temp_files)

if __name__ == "__main__":
    docx_path = "temp\\temp_document_pc.docx"  # Ruta del archivo DOCX
    pdf_path = convertir_a_pdf(docx_path)  # Convertir a PDF
    if pdf_path:
        mostrar_pdf(pdf_path, docx_path)  # Mostrar el PDF
